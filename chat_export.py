"""
Chat Export Module - Generate PDF/DOCX from chat messages with markdown support.
Supports dark and light themes for both output formats.
"""
import re
from io import BytesIO
from datetime import datetime


def generate_pdf(messages, theme='dark'):
    """
    Generate PDF from chat messages.
    Uses pre-rendered HTML (with SVGs converted to PNG) if available,
    otherwise falls back to markdown rendering.

    Args:
        messages: list of dicts with keys: role, text, timestamp, rendered_html (optional)
        theme: 'dark' or 'light'

    Returns:
        BytesIO buffer containing the PDF
    """
    import markdown
    from xhtml2pdf import pisa

    md = markdown.Markdown(extensions=['tables', 'fenced_code'])

    css = _get_pdf_css(theme)
    html_parts = [f'<html><head><meta charset="utf-8"><style>{css}</style></head><body>']
    html_parts.append('<h1 class="export-title">Chat Export</h1>')
    html_parts.append(f'<p class="export-meta">Exported: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")} | {len(messages)} message(s)</p>')
    html_parts.append('<hr>')

    for msg in messages:
        role_class = 'user' if msg['role'] == 'user' else 'agent'
        label = 'You' if msg['role'] == 'user' else 'Agent'
        timestamp = msg.get('timestamp', '')

        rendered_html = msg.get('rendered_html')
        if rendered_html:
            # Use the pre-rendered HTML from the browser (includes Mermaid PNGs, tables, etc.)
            content_html = _clean_rendered_html(rendered_html, theme=theme)
        else:
            # Fallback: render raw markdown
            md.reset()
            content_html = md.convert(msg.get('text', ''))

        html_parts.append(f'''
        <div class="message {role_class}">
            <div class="msg-header">
                <span class="msg-role">{label}</span>
                <span class="msg-time">{timestamp}</span>
            </div>
            <div class="msg-content">{content_html}</div>
        </div>
        ''')

    html_parts.append('</body></html>')
    full_html = '\n'.join(html_parts)

    buffer = BytesIO()
    pisa.CreatePDF(full_html, dest=buffer)
    buffer.seek(0)
    return buffer


def generate_docx(messages, theme='dark'):
    """
    Generate DOCX from chat messages with markdown formatting.

    Args:
        messages: list of dicts with keys: role, text, timestamp
        theme: 'dark' or 'light'

    Returns:
        BytesIO buffer containing the DOCX
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    colors = _get_docx_colors(theme)

    if theme == 'dark':
        _set_docx_page_background(doc, colors['page_bg'])

    # Title
    title_para = doc.add_paragraph()
    run = title_para.add_run('Chat Export')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = colors['title']

    # Metadata
    meta_para = doc.add_paragraph()
    run = meta_para.add_run(f'Exported: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")} | {len(messages)} message(s)')
    run.font.size = Pt(10)
    run.font.color.rgb = colors['meta']

    doc.add_paragraph()  # spacer

    for msg in messages:
        label = 'You' if msg['role'] == 'user' else 'Agent'
        role_color = colors['user_label'] if msg['role'] == 'user' else colors['agent_label']
        timestamp = msg.get('timestamp', '')

        # Role + timestamp header
        header_para = doc.add_paragraph()
        role_run = header_para.add_run(label)
        role_run.bold = True
        role_run.font.size = Pt(11)
        role_run.font.color.rgb = role_color

        if timestamp:
            ts_run = header_para.add_run(f'  {timestamp}')
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = colors['timestamp']

        # Message content: markdown text formatting + embedded PNG images from rendered HTML
        _add_markdown_to_docx(doc, msg.get('text', ''), colors)

        # Extract and insert PNG images (Mermaid diagrams, LaTeX formulas) from rendered HTML
        rendered_html = msg.get('rendered_html')
        if rendered_html:
            _insert_images_from_html(doc, rendered_html, theme=theme)

        doc.add_paragraph()  # spacer between messages

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def _get_pdf_font_face():
    """Return @font-face CSS for Unicode-capable DejaVu Sans font."""
    return '''
            @font-face {
                font-family: 'DejaVu';
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf');
            }
            @font-face {
                font-family: 'DejaVu';
                font-weight: bold;
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf');
            }
            @font-face {
                font-family: 'DejaVuMono';
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf');
            }
            @font-face {
                font-family: 'DejaVuMono';
                font-weight: bold;
                src: url('/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf');
            }
    '''


def _get_pdf_css(theme):
    """Return CSS string for PDF based on theme."""
    font_face = _get_pdf_font_face()
    if theme == 'dark':
        return font_face + '''
            @page { size: A4; margin: 1.5cm; background-color: #121212; }
            body { background-color: #121212; color: #e0e0e0; font-family: 'DejaVu', sans-serif; font-size: 11px; }
            .export-title { color: #7aa5d2; font-size: 20px; margin-bottom: 2px; }
            .export-meta { color: #888; font-size: 10px; margin-top: 0; }
            hr { border: none; border-top: 1px solid #444; }
            .message { margin-bottom: 14px; padding: 10px; border-radius: 6px; }
            .message.user { background-color: #0e639c; }
            .message.agent { background-color: #333333; }
            .msg-header { margin-bottom: 5px; }
            .msg-role { font-weight: bold; font-size: 11px; }
            .msg-time { font-size: 9px; color: #aaa; margin-left: 10px; }
            .msg-content { line-height: 1.5; }
            .msg-content code { background-color: rgba(0,0,0,0.3); padding: 1px 4px; font-family: 'DejaVuMono', monospace; font-size: 10px; }
            .msg-content pre { background-color: rgba(0,0,0,0.35); padding: 8px; font-family: 'DejaVuMono', monospace; font-size: 10px; }
            .msg-content table { border-collapse: collapse; width: 100%; margin: 5px 0; }
            .msg-content th, .msg-content td { border: 1px solid #555; padding: 4px 8px; font-size: 10px; }
            .msg-content th { background-color: rgba(255,255,255,0.08); }
            .msg-content h1, .msg-content h2, .msg-content h3 { color: #7aa5d2; }
            .msg-content strong { color: #ffffff; }
            .msg-content blockquote { border-left: 3px solid #555; padding-left: 10px; color: #aaa; }
            .msg-content ul, .msg-content ol { padding-left: 20px; }
            .msg-content img { max-width: 100%; height: auto; margin: 6px 0; }
            .mermaid-wrapper { display: block; }
            .diff-add { color: #4ec9b0; display: block; }
            .diff-del { color: #f48771; display: block; }
            .diff-hunk { color: #9cdcfe; display: block; }
            .diff-ctx { display: block; }
        '''
    else:
        return font_face + '''
            @page { size: A4; margin: 1.5cm; }
            body { background-color: #ffffff; color: #222222; font-family: 'DejaVu', sans-serif; font-size: 11px; }
            .export-title { color: #2c3e50; font-size: 20px; margin-bottom: 2px; }
            .export-meta { color: #888; font-size: 10px; margin-top: 0; }
            hr { border: none; border-top: 1px solid #ddd; }
            .message { margin-bottom: 14px; padding: 10px; border-radius: 6px; border: 1px solid #ddd; }
            .message.user { background-color: #e8f4fd; }
            .message.agent { background-color: #f5f5f5; }
            .msg-header { margin-bottom: 5px; }
            .msg-role { font-weight: bold; font-size: 11px; color: #333; }
            .msg-time { font-size: 9px; color: #999; margin-left: 10px; }
            .msg-content { line-height: 1.5; }
            .msg-content code { background-color: #f0f0f0; padding: 1px 4px; font-family: 'DejaVuMono', monospace; font-size: 10px; }
            .msg-content pre { background-color: #f5f5f5; padding: 8px; font-family: 'DejaVuMono', monospace; font-size: 10px; border: 1px solid #ddd; }
            .msg-content table { border-collapse: collapse; width: 100%; margin: 5px 0; }
            .msg-content th, .msg-content td { border: 1px solid #ccc; padding: 4px 8px; font-size: 10px; }
            .msg-content th { background-color: #f0f0f0; }
            .msg-content h1, .msg-content h2, .msg-content h3 { color: #2c3e50; }
            .msg-content blockquote { border-left: 3px solid #ddd; padding-left: 10px; color: #666; }
            .msg-content ul, .msg-content ol { padding-left: 20px; }
            .msg-content img { max-width: 100%; height: auto; margin: 6px 0; }
            .mermaid-wrapper { display: block; }
            .diff-add { color: #1a7340; display: block; }
            .diff-del { color: #c0392b; display: block; }
            .diff-hunk { color: #2471a3; display: block; }
            .diff-ctx { display: block; }
        '''


# ---------------------------------------------------------------------------
# LaTeX rendering via matplotlib mathtext
# ---------------------------------------------------------------------------

def _render_latex_matplotlib(latex, is_display=False, theme='dark'):
    """
    Render a LaTeX math string to a base64 PNG using matplotlib's mathtext engine.
    Returns base64-encoded PNG string, or None on failure.
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import io, base64 as _b64

        text_color = '#d4d4d4' if theme == 'dark' else '#1a1a1a'
        fontsize = 14 if is_display else 11
        render_dpi = 150

        fig = plt.figure(figsize=(0.1, 0.1), dpi=render_dpi, facecolor='none')
        text_obj = fig.text(0, 0, f'${latex}$', fontsize=fontsize,
                            color=text_color, ha='left', va='bottom')

        # Draw once to get actual pixel bbox, then resize figure to fit
        fig.canvas.draw()
        renderer = fig.canvas.get_renderer()
        bbox = text_obj.get_window_extent(renderer=renderer)

        pad_px = 8
        fig.set_size_inches(
            max((bbox.width + pad_px * 2) / render_dpi, 0.5),
            max((bbox.height + pad_px * 2) / render_dpi, 0.25)
        )

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=render_dpi, bbox_inches='tight',
                    transparent=True, facecolor='none', pad_inches=pad_px / render_dpi)
        plt.close(fig)
        buf.seek(0)
        return _b64.b64encode(buf.read()).decode('ascii')
    except Exception as e:
        print(f'[LaTeX render] {e}')
        return None


# ---------------------------------------------------------------------------
# Rendered HTML helpers (for pre-rendered content from browser)
# ---------------------------------------------------------------------------

def _clean_rendered_html(html, theme='dark'):
    """
    Clean browser-rendered HTML for PDF embedding.
    - Removes UI-only elements and dangerous attributes.
    - Renders LaTeX block formulas to PNG via matplotlib.
    - Converts <img src="data:image/svg+xml;base64,..."> to PNG via cairosvg.
    """
    import base64 as _b64

    # Remove script tags
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
    # Remove event attributes
    html = re.sub(r'\s+on\w+="[^"]*"', '', html)

    # Render LaTeX formulas (placed by captureRenderedHtml as <img class="latex-png">)
    # Block/display formulas → matplotlib PNG (renders perfectly as block element)
    # Inline formulas → styled italic span (xhtml2pdf can't flow <img> inline in text)
    def latex_img_to_png(m):
        try:
            attrs = m.group(0)
            b64_match = re.search(r'data-latex-b64="([^"]+)"', attrs)
            display_match = re.search(r'data-display="([^"]*)"', attrs)
            if not b64_match:
                return m.group(0)
            latex = _b64.b64decode(b64_match.group(1) + '==').decode('utf-8')
            is_display = display_match and display_match.group(1) == '1'
            if is_display:
                # Block formula: render as PNG via matplotlib
                png_b64 = _render_latex_matplotlib(latex, is_display=True, theme=theme)
                if png_b64:
                    return (f'<p style="text-align:center;margin:6px 0;">'
                            f'<img src="data:image/png;base64,{png_b64}" '
                            f'style="max-width:100%;height:auto;"/></p>')
            # Inline formula: styled italic text — xhtml2pdf can't flow images inline
            escaped = latex.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            return f'<i style="font-family:serif;">{escaped}</i>'
        except Exception as e:
            return m.group(0)
    html = re.sub(r'<img\b[^>]*class="latex-png"[^>]*/?>',
                  latex_img_to_png, html, flags=re.IGNORECASE)

    # xhtml2pdf ignores white-space:pre-wrap on literal \n — convert to <br/> inside <pre> blocks
    def fix_pre_newlines(m):
        return '<pre>' + m.group(1).replace('\n', '<br/>') + '</pre>'
    html = re.sub(r'<pre>(.*?)</pre>', fix_pre_newlines, html, flags=re.DOTALL | re.IGNORECASE)

    # Convert SVG data URIs to PNG data URIs using cairosvg
    def svg_to_png_img(m):
        try:
            import cairosvg
            svg_bytes = _b64.b64decode(m.group(1).strip())
            png_bytes = cairosvg.svg2png(bytestring=svg_bytes, dpi=200, output_width=1600)
            png_b64 = _b64.b64encode(png_bytes).decode('ascii')
            return (f'<img src="data:image/png;base64,{png_b64}" '
                    f'style="max-width:100%;height:auto;margin:8px 0;display:block;" />')
        except Exception as e:
            return f'<p style="color:#888;font-style:italic">[Diagram: {e}]</p>'

    html = re.sub(
        r'<img[^>]+src="data:image/svg\+xml;base64,([^"]+)"[^>]*/?>',
        svg_to_png_img,
        html
    )
    return html


def _insert_images_from_html(doc, rendered_html, theme='dark'):
    """
    Parse rendered HTML and insert images (Mermaid SVGs, LaTeX display formulas) into DOCX.
    """
    import base64 as _b64
    from docx.shared import Inches, Pt, RGBColor

    # 1. Insert PNG images (LaTeX display formulas captured from browser canvas)
    for b64data in re.findall(
            r'<img[^>]+src="data:image/png;base64,([^"]+)"[^>]*/?>',
            rendered_html):
        try:
            png_bytes = _b64.b64decode(b64data.strip())
            doc.add_picture(BytesIO(png_bytes), width=Inches(5))
        except Exception:
            pass

    # 2. Insert SVG diagrams (Mermaid) via cairosvg → PNG
    for b64data in re.findall(
            r'<img[^>]+src="data:image/svg\+xml;base64,([^"]+)"[^>]*/?>',
            rendered_html):
        try:
            svg_bytes = _b64.b64decode(b64data.strip())
            try:
                import cairosvg
                png_bytes = cairosvg.svg2png(bytestring=svg_bytes, dpi=200, output_width=1600)
                doc.add_picture(BytesIO(png_bytes), width=Inches(6))
                continue
            except ImportError:
                pass
            p = doc.add_paragraph()
            r = p.add_run('[Diagram — export as PDF to view]')
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _get_docx_colors(theme):
    """Return color dict for DOCX based on theme."""
    from docx.shared import RGBColor
    if theme == 'dark':
        return {
            'page_bg': '121212',
            'title': RGBColor(0x7a, 0xa5, 0xd2),
            'meta': RGBColor(0x88, 0x88, 0x88),
            'text': RGBColor(0xe0, 0xe0, 0xe0),
            'user_label': RGBColor(0x4C, 0xAF, 0x50),
            'agent_label': RGBColor(0x21, 0x96, 0xF3),
            'timestamp': RGBColor(0xaa, 0xaa, 0xaa),
            'heading': RGBColor(0x7a, 0xa5, 0xd2),
            'bold': RGBColor(0xff, 0xff, 0xff),
            'code': RGBColor(0xce, 0x91, 0x78),
            'link': RGBColor(0x7a, 0xa5, 0xd2),
        }
    else:
        return {
            'page_bg': 'ffffff',
            'title': RGBColor(0x2c, 0x3e, 0x50),
            'meta': RGBColor(0x88, 0x88, 0x88),
            'text': RGBColor(0x22, 0x22, 0x22),
            'user_label': RGBColor(0x2e, 0x7d, 0x32),
            'agent_label': RGBColor(0x15, 0x65, 0xc0),
            'timestamp': RGBColor(0x99, 0x99, 0x99),
            'heading': RGBColor(0x2c, 0x3e, 0x50),
            'bold': RGBColor(0x00, 0x00, 0x00),
            'code': RGBColor(0xc7, 0x25, 0x4e),
            'link': RGBColor(0x15, 0x65, 0xc0),
        }


def _set_docx_page_background(doc, hex_color):
    """Set page background color for DOCX document."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    bg_elem = OxmlElement('w:background')
    bg_elem.set(qn('w:color'), hex_color)
    bg_elem.set(qn('w:themeColor'), 'dark1')
    doc.element.insert(0, bg_elem)


def _add_markdown_to_docx(doc, text, colors):
    """
    Parse markdown text and add formatted paragraphs to DOCX.
    Handles: headings, bold, italic, code inline, code blocks, lists, links.
    """
    from docx.shared import Pt

    lines = text.split('\n')
    in_code_block = False
    code_block_lines = []

    for line in lines:
        # Code block start/end
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_para = doc.add_paragraph()
                code_text = '\n'.join(code_block_lines)
                code_run = code_para.add_run(code_text)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = colors['code']
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # Headings
        header_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if header_match:
            level = len(header_match.group(1))
            para = doc.add_paragraph()
            run = para.add_run(header_match.group(2))
            run.bold = True
            run.font.size = Pt(max(10, 18 - (level * 2)))
            run.font.color.rgb = colors['heading']
            continue

        # List items (unordered and ordered)
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)', line)
        if list_match:
            indent = len(list_match.group(1))
            content = list_match.group(3)
            bullet = list_match.group(2)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(20 + indent * 10)
            # Add bullet/number prefix
            prefix = '\u2022 ' if bullet in ('-', '*', '+') else f'{bullet} '
            prefix_run = para.add_run(prefix)
            prefix_run.font.size = Pt(10)
            prefix_run.font.color.rgb = colors['text']
            _add_inline_formatting(para, content, colors)
            continue

        # Empty lines - skip
        if not line.strip():
            continue

        # Regular paragraph with inline formatting
        para = doc.add_paragraph()
        _add_inline_formatting(para, line, colors)


def _add_inline_formatting(paragraph, text, colors):
    """
    Add inline markdown formatting to a paragraph.
    Handles: **bold**, *italic*, `code`, [text](url)
    """
    from docx.shared import Pt

    # Split text by inline markdown patterns
    # Order matters: **bold** before *italic*
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = colors.get('bold', colors['text'])
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.color.rgb = colors['text']
            run.font.size = Pt(10)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = colors['code']
        elif part.startswith('['):
            link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if link_match:
                run = paragraph.add_run(link_match.group(1))
                run.font.color.rgb = colors['link']
                run.underline = True
                run.font.size = Pt(10)
            else:
                run = paragraph.add_run(part)
                run.font.color.rgb = colors['text']
                run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.color.rgb = colors['text']
            run.font.size = Pt(10)
